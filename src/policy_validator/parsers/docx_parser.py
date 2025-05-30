"""Microsoft Word document parser for policy validation.

This module provides functionality for extracting and analyzing content from
Microsoft Word (.docx) documents using the python-docx library. It extracts
text, headings, and structural information needed for policy validation.

Example:
    parser = DocxParser("policy.docx")
    content = parser.parse()
    text = content['text']
    headings = content['headings']

Note:
    Only supports .docx format (newer Word documents). For .doc files,
    convert to .docx first. Cannot process password-protected files.
"""

import docx


class DocxParser:
    """Parser for .docx documents using python-docx.
    
    This class handles the extraction of content from .docx files,
    organizing it into a structured format for validation.

    Attributes:
        file_path (str): Path to the .docx file
        document: python-docx Document object

    Parsing Capabilities:
        - Full text extraction
        - Heading level detection
        - Paragraph count
        - Basic structure analysis

    Error Handling:
        - Invalid file format
        - Corrupted documents
        - Access permission issues
    """
    
    def __init__(self, file_path):
        """
        Initialize the DocxParser.
        
        Args:
            file_path (str): Path to the .docx file
        """
        self.file_path = file_path
        self.document = None
        
    def parse(self):
        """Parse the document and return its content.
        
        Opens and processes the .docx file, extracting its content into a
        structured dictionary format suitable for policy validation.
        
        Returns:
            dict: Document content containing:
                text (str): Full plain text of the document
                headings (list): List of dictionaries with heading level and text
                paragraphs (int): Total number of paragraphs
                
        Raises:
            FileNotFoundError: If the document file cannot be found
            docx.exceptions.PackageNotFoundError: If file is not a valid .docx
            PermissionError: If file cannot be accessed due to permissions
        """
        self.document = docx.Document(self.file_path)
        
        content = {
            'text': self._extract_text(),
            'headings': self._extract_headings(),
            'paragraphs': len(self.document.paragraphs),
        }
        
        return content
    
    def _extract_text(self):
        """Extract full text from the document.
        
        Concatenates all paragraph text from the document into a single string,
        preserving paragraph breaks with newlines.
        
        Returns:
            str: Full text content of the document
            
        Note:
            This method ignores text in tables, headers, footers, and text boxes.
            Only paragraph text is extracted.
        """
        return '\n'.join([para.text for para in self.document.paragraphs])
    
    def _extract_headings(self):
        """Extract headings from the document.
        
        Identifies paragraphs with Heading styles and extracts their level and text.
        Heading levels are determined by the style name (e.g., 'Heading 1', 'Heading 2').
        
        Returns:
            list: List of dictionaries containing:
                level (int): Heading level (1 for Heading 1, etc.)
                text (str): Heading text content
                
        Note:
            Only paragraphs with standard Word heading styles are detected.
            Custom styles or manually formatted headings may not be identified.
        """
        headings = []
        
        for para in self.document.paragraphs:
            if para.style.name.startswith('Heading'):
                headings.append({
                    'level': int(para.style.name.replace('Heading ', '')),
                    'text': para.text
                })
                
        return headings

